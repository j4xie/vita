"""
文章导出器 - 支持 JSON / DOCX 格式
"""
import json
import os
import re
import logging
from typing import List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .parser import ArticleData

logger = logging.getLogger(__name__)


def safe_filename(name: str, max_len: int = 80) -> str:
    """生成安全的文件名"""
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', name)
    return name[:max_len].strip('_ ')


def export_json(article: ArticleData, output_dir: str) -> str:
    """导出为 JSON 文件"""
    os.makedirs(output_dir, exist_ok=True)
    filename = safe_filename(article.title or 'untitled') + '.json'
    filepath = os.path.join(output_dir, filename)

    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(article.to_dict(), f, ensure_ascii=False, indent=2)

    logger.info(f'[Export] JSON saved: {filepath}')
    return filepath


def export_docx(article: ArticleData, output_dir: str) -> str:
    """导出为 DOCX 文件"""
    os.makedirs(output_dir, exist_ok=True)
    filename = safe_filename(article.title or 'untitled') + '.docx'
    filepath = os.path.join(output_dir, filename)

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)

    # 标题
    heading = doc.add_heading(article.title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息行
    meta_parts = []
    if article.account_name:
        meta_parts.append(f'公众号: {article.account_name}')
    if article.author:
        meta_parts.append(f'作者: {article.author}')
    if article.publish_date:
        meta_parts.append(f'发布: {article.publish_date}')

    if meta_parts:
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta_para.add_run(' | '.join(meta_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # 分割线
    doc.add_paragraph('─' * 50)

    # 正文 - 按段落处理
    if article.content_text:
        paragraphs = article.content_text.split('\n')
        for para_text in paragraphs:
            text = para_text.strip()
            if not text:
                continue
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

    # 来源链接
    if article.source_url:
        doc.add_paragraph('')
        doc.add_paragraph('─' * 50)
        source_para = doc.add_paragraph()
        run = source_para.add_run(f'原文链接: {article.source_url}')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(filepath)
    logger.info(f'[Export] DOCX saved: {filepath}')
    return filepath


def export_batch_index(articles: List[ArticleData], output_dir: str) -> str:
    """生成批量爬取的 index.json 索引"""
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, 'index.json')

    index = []
    for a in articles:
        index.append({
            'title': a.title,
            'account_name': a.account_name,
            'publish_date': a.publish_date,
            'source_url': a.source_url,
            'content_length': len(a.content_text),
            'image_count': len(a.images),
        })

    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(index, f, ensure_ascii=False, indent=2)

    logger.info(f'[Export] Index saved: {filepath} ({len(articles)} articles)')
    return filepath
