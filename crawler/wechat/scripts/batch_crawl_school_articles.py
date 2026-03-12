"""
批量爬取 school_data/各分部新生推文攻略.docx 中的所有微信公众号文章
按学校分类存储为纯文本DOCX到 school_data/{学校名}/ 目录
"""
import asyncio
import logging
import os
import re
import sys
import time

# 项目路径设置
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
CRAWLER_ROOT = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, CRAWLER_ROOT)

from src.official_account.crawler import ArticleCrawler
from src.official_account.exporter import safe_filename

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 只写文件日志，终端输出用 print 可视化
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(SCRIPT_DIR, 'batch_crawl.log'), encoding='utf-8'),
    ]
)
logger = logging.getLogger(__name__)

# school_data 目录
SCHOOL_DATA_DIR = os.path.abspath(os.path.join(CRAWLER_ROOT, '..', '..', '..', 'school_data'))
DOCX_FILE = os.path.join(SCHOOL_DATA_DIR, '各分部新生推文攻略.docx')

# 学校名 -> school_data文件夹名 映射
SCHOOL_FOLDER_MAP = {
    'ChineseUnionUCSC': 'UCSC',
    'CU at UCD': 'UCD',
    'CU UCSB': 'UCSB',
    'CU UCSD': 'UCSD',
    'CU UMN': 'UMN',
    'CU USC': 'USC',
    'CU UW': 'UW',
    'CU在圣克鲁兹': 'UCSC',
    'UCB CU': 'UCB',
    'UCI CU': 'UCI',
    'UCLA CU': 'UCLA',
}

SCHOOL_NAMES = set(SCHOOL_FOLDER_MAP.keys())

# 暂时跳过的学校 (如公众号迁移导致无法爬取)
SKIP_SCHOOLS = {'UCD'}

# ─── 可视化工具 ───────────────────────────────────────────

COLOR_RESET = '\033[0m'
COLOR_GREEN = '\033[92m'
COLOR_RED = '\033[91m'
COLOR_YELLOW = '\033[93m'
COLOR_CYAN = '\033[96m'
COLOR_BOLD = '\033[1m'
COLOR_DIM = '\033[2m'


def progress_bar(current, total, width=30):
    """生成进度条字符串"""
    filled = int(width * current / total) if total else 0
    bar = '█' * filled + '░' * (width - filled)
    pct = current / total * 100 if total else 0
    return f'[{bar}] {pct:5.1f}%'


def format_time(seconds):
    """格式化秒数为 mm:ss"""
    m, s = divmod(int(seconds), 60)
    return f'{m:02d}:{s:02d}'


def print_header():
    print(f'''
{COLOR_BOLD}╔══════════════════════════════════════════════════════════╗
║        微信公众号文章批量爬取工具 v1.0                   ║
║        WeChat Article Batch Crawler                      ║
╚══════════════════════════════════════════════════════════╝{COLOR_RESET}
''')


def print_school_table(schools):
    """打印学校统计表格"""
    print(f'{COLOR_BOLD}┌──────────┬────────┐')
    print(f'│  学校    │ 文章数 │')
    print(f'├──────────┼────────┤{COLOR_RESET}')
    total = 0
    for school, articles in schools.items():
        count = len(articles)
        total += count
        # 补齐中文宽度
        name = school.ljust(8)
        print(f'│  {COLOR_CYAN}{name}{COLOR_RESET}│  {count:>4}  │')
    print(f'{COLOR_BOLD}├──────────┼────────┤')
    print(f'│  总计    │  {total:>4}  │')
    print(f'└──────────┴────────┘{COLOR_RESET}')
    print()
    return total


def print_result(success, fail, skip, total, elapsed):
    """打印最终结果"""
    print(f'''
{COLOR_BOLD}╔══════════════════════════════════════════════════════════╗
║                      爬取完成!                           ║
╠══════════════════════════════════════════════════════════╣{COLOR_RESET}
  {COLOR_GREEN}✓ 成功: {success}{COLOR_RESET}
  {COLOR_RED}✗ 失败: {fail}{COLOR_RESET}
  {COLOR_YELLOW}⊘ 跳过: {skip}{COLOR_RESET}
  {COLOR_DIM}  总计: {total}  |  耗时: {format_time(elapsed)}{COLOR_RESET}
{COLOR_BOLD}╚══════════════════════════════════════════════════════════╝{COLOR_RESET}''')


# ─── 解析 & 导出 ─────────────────────────────────────────

def parse_docx_file(filepath: str) -> dict:
    """解析DOCX文件，返回 {学校文件夹名: [(文章标题, URL), ...]}"""
    doc = DocxDocument(filepath)
    schools = {}
    current_school = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text in SCHOOL_NAMES and 'mp.weixin.qq.com' not in text:
            folder = SCHOOL_FOLDER_MAP[text]
            current_school = folder
            if current_school not in schools:
                schools[current_school] = []
            continue

        if current_school and 'mp.weixin.qq.com' in text:
            url_match = re.search(r'https?://mp\.weixin\.qq\.com/s/[A-Za-z0-9_-]+', text)
            if url_match:
                url = url_match.group(0)
            else:
                bare_match = re.search(r'mp\.weixin\.qq\.com/s/[A-Za-z0-9_-]+', text)
                if bare_match:
                    url = 'https://' + bare_match.group(0)
                else:
                    continue

            url = url.replace('http://', 'https://')
            title = re.split(r'https?://mp\.weixin\.qq\.com', text)[0].strip()
            if not title:
                title = re.split(r'mp\.weixin\.qq\.com', text)[0].strip()

            schools[current_school].append((title, url))

    return schools


def export_plain_text_docx(article, output_dir: str) -> str:
    """导出纯文本DOCX"""
    os.makedirs(output_dir, exist_ok=True)
    filename = safe_filename(article.title or 'untitled') + '.docx'
    filepath = os.path.join(output_dir, filename)

    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)

    heading = doc.add_heading(article.title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if article.content_text:
        for para_text in article.content_text.split('\n'):
            text = para_text.strip()
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)

    doc.save(filepath)
    return filepath


# ─── 主流程 ──────────────────────────────────────────────

async def main():
    print_header()

    print(f'{COLOR_DIM}读取: {DOCX_FILE}{COLOR_RESET}')
    schools = parse_docx_file(DOCX_FILE)

    # 跳过迁移学校
    if SKIP_SCHOOLS:
        skipped = {k: v for k, v in schools.items() if k in SKIP_SCHOOLS}
        schools = {k: v for k, v in schools.items() if k not in SKIP_SCHOOLS}
        for s, a in skipped.items():
            print(f'{COLOR_YELLOW}⊘ 跳过学校: {s} ({len(a)} 篇) — 公众号已迁移{COLOR_RESET}')

    total_articles = print_school_table(schools)

    crawler = ArticleCrawler(timeout=30.0)
    success_count = 0
    fail_count = 0
    skip_count = 0
    migrate_count = 0
    migrated_articles = []  # 记录迁移失败的文章
    global_idx = 0
    start_time = time.time()

    try:
        for school, articles in schools.items():
            school_dir = os.path.join(SCHOOL_DATA_DIR, school)
            os.makedirs(school_dir, exist_ok=True)

            print(f'\n{COLOR_BOLD}── {school} ({len(articles)} 篇) ──{COLOR_RESET}')
            print(f'{COLOR_DIM}   目录: {school_dir}{COLOR_RESET}')

            for i, (title, url) in enumerate(articles, 1):
                global_idx += 1
                elapsed = time.time() - start_time
                eta = (elapsed / global_idx) * (total_articles - global_idx) if global_idx else 0

                # 进度条
                bar = progress_bar(global_idx, total_articles)
                status_line = f'  {bar}  {global_idx}/{total_articles}  ETA {format_time(eta)}'

                # 检查是否已存在
                expected_filename = safe_filename(title or 'untitled') + '.docx'
                expected_path = os.path.join(school_dir, expected_filename)
                if os.path.exists(expected_path):
                    print(f'{status_line}  {COLOR_YELLOW}⊘ 跳过{COLOR_RESET} {title[:40]}')
                    logger.info(f'[{school}] 跳过: {expected_filename}')
                    skip_count += 1
                    continue

                display_title = title[:40] if title else '(无标题)'
                print(f'{status_line}  {COLOR_CYAN}↓ 爬取{COLOR_RESET} {display_title}', end='', flush=True)
                logger.info(f'[{school}] 爬取: {title} | {url}')

                try:
                    article = await crawler.crawl_article(url)
                    if article and article.content_text:
                        # 检查是否爬到的是迁移页/验证码页 (内容很短且无实质内容)
                        if len(article.content_text) < 50 and ('迁移' in article.content_text or '验证' in article.content_text):
                            print(f'\r{status_line}  {COLOR_YELLOW}⇄ 迁移{COLOR_RESET} {display_title}')
                            logger.warning(f'[{school}] 迁移: {title} | {url}')
                            migrated_articles.append((school, title, url))
                            migrate_count += 1
                        else:
                            filepath = export_plain_text_docx(article, school_dir)
                            chars = len(article.content_text)
                            print(f'\r{status_line}  {COLOR_GREEN}✓ 成功{COLOR_RESET} {article.title[:40]} ({chars}字)')
                            logger.info(f'[{school}] 成功: {os.path.basename(filepath)} ({chars}字)')
                            success_count += 1
                    else:
                        # 检查是否为迁移导致的空内容
                        import httpx as _httpx
                        _client = await crawler._get_client()
                        _resp = await _client.get(url)
                        if crawler._is_migrated(_resp.text):
                            print(f'\r{status_line}  {COLOR_YELLOW}⇄ 迁移{COLOR_RESET} {display_title}')
                            logger.warning(f'[{school}] 迁移: {title} | {url}')
                            migrated_articles.append((school, title, url))
                            migrate_count += 1
                        else:
                            print(f'\r{status_line}  {COLOR_RED}✗ 空内容{COLOR_RESET} {display_title}')
                            logger.warning(f'[{school}] 失败(空内容): {title} | {url}')
                            fail_count += 1
                except Exception as e:
                    print(f'\r{status_line}  {COLOR_RED}✗ 错误{COLOR_RESET} {display_title} | {str(e)[:30]}')
                    logger.error(f'[{school}] 错误: {title} | {url} | {e}')
                    fail_count += 1

                # 间隔避免被封
                await asyncio.sleep(2.0)

    except KeyboardInterrupt:
        print(f'\n\n{COLOR_YELLOW}⚠ 用户中断 (Ctrl+C){COLOR_RESET}')
    finally:
        await crawler.close()

    elapsed = time.time() - start_time
    print_result(success_count, fail_count, skip_count, total_articles, elapsed)

    # 输出迁移文章列表
    if migrated_articles:
        migrate_file = os.path.join(SCHOOL_DATA_DIR, 'migrated_articles.txt')
        with open(migrate_file, 'w', encoding='utf-8') as f:
            f.write('# 以下文章因公众号迁移无法自动爬取，需要手动更新链接\n')
            f.write(f'# 共 {len(migrated_articles)} 篇\n\n')
            current_s = ''
            for school, title, url in migrated_articles:
                if school != current_s:
                    f.write(f'\n## {school}\n')
                    current_s = school
                f.write(f'  {title}\n  {url}\n\n')

        print(f'\n{COLOR_YELLOW}⚠ 有 {len(migrated_articles)} 篇文章因公众号迁移无法自动爬取{COLOR_RESET}')
        print(f'{COLOR_DIM}  迁移列表已保存到: {migrate_file}{COLOR_RESET}')
        print(f'{COLOR_DIM}  这些文章需要你手动在浏览器中打开原链接 -> 点击"访问文章" -> 复制新链接{COLOR_RESET}')


if __name__ == '__main__':
    # Windows终端UTF-8支持
    if sys.platform == 'win32':
        sys.stdout.reconfigure(encoding='utf-8')
    asyncio.run(main())
